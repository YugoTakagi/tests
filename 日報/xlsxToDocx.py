import docx
import xlrd
import pprint
import xlwt
from docx.shared import RGBColor
from docx.shared import Pt
from tqdm import tqdm

#####__def__####################################################################
def makingNameList(col):
    name = [""]
    # if "名前"==str(col[0]) or "なまえ"==str(col[0]) or "name"== str(col[0]):
    for ii, ee in enumerate(col):
        if ii == 0:
            continue
        else:
            name.append(col[ii])
    # print("makingNameList {}".format(col))
    return name
################################################################################
#####__init__###################################
sheetList = input("ファイル名(xxx)を入力(xxx.xlsx形式)>>> ")
sheetNames = sheetList.split(",")

for sheetName in tqdm(sheetNames):
    doc = docx.Document()
    wb = xlrd.open_workbook(sheetName +".xlsx")
    sheet = wb.sheet_by_index(0)
    #####
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    # paragraph.size = Pt(16)
    paragraph.text = sheetName
     #####__footer__#################################
    section2 = doc.sections[0]
    footers2 = section2.footer  # a HeadersFooters collection object
    default_footer = footers2.paragraphs[0]
    default_footer.text = "kogakuin   robot   project"
    ################################################
    ################################################
    row = sheet.row_values(0)#質問内容の保存
    cols = []
    name = []
    for i, e in enumerate(row):
        col = sheet.col_values(i)
        cols.append(col)
        # print("col = {}".format(col[0]))
        #####__making_name__####################################################
        if "名前"==str(col[0]) or "なまえ"==str(col[0]) or "name"== str(col[0]):
            name = makingNameList(col)#名前の配列の作成
        ########################################################################

    for i, e in enumerate(row):
        # if i == 0:
        #     continue
        if "名前" == e or "なまえ" == e or "name" == e:
            continue
        elif "タイムスタンプ" == e:
            continue
        else:
            for ii, ee in enumerate(cols[i]):
                # print("{}>>> {}".format(i, cols[i][ii]))
                if ii == 0:
                    doc.add_heading(cols[i][ii], 0)
                else:
                    if str(cols[i][ii]) == "":
                        continue
                    else:
                        run = doc.add_paragraph().add_run(str(name[ii]) +"  :")
                        font = run.font
                        font.color.rgb = RGBColor(47, 145, 79)
                        font.size = Pt(8)
                        #####
                        doc.add_paragraph("   " +str(cols[i][ii]))
    doc.add_heading("Thenk you very much today.", 0)
    for i, e in enumerate(name):
        if i == 0:
            continue
        else:
            doc.add_paragraph("   " +e)
    ################################################
    paragraph.text = sheetName +" #出席人数:{}人".format(len(name)-1) +"\t\tKRP"
    doc.save('./{}.docx'.format(sheetName))
print("fin")
